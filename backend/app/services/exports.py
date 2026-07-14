from collections.abc import Mapping, Sequence
from datetime import datetime
from io import BytesIO
from zoneinfo import ZoneInfo

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Cm, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill

MOSCOW = ZoneInfo("Europe/Moscow")
HEADERS = ("Группа", "VK ID", "Студент", "Статус", "Ответ", "Время ответа", "Изображения")


def create_xlsx(title: str, results: Sequence[Mapping[str, object]]) -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Результаты"
    sheet.append((f"Рассылка: {title}",))
    sheet.append(HEADERS)
    for result in results:
        sheet.append(_row(result))

    sheet.freeze_panes = "A3"
    sheet.auto_filter.ref = f"A2:G{max(sheet.max_row, 2)}"
    sheet["A1"].font = Font(name="Arial", size=14, bold=True)
    for cell in sheet[2]:
        cell.font = Font(name="Arial", bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="2F5597")
    for row in sheet.iter_rows(min_row=3):
        for cell in row:
            cell.font = Font(name="Arial", size=10)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    for column, width in zip("ABCDEFG", (20, 14, 28, 22, 40, 22, 45), strict=True):
        sheet.column_dimensions[column].width = width

    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def create_docx(title: str, results: Sequence[Mapping[str, object]]) -> bytes:
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = Cm(29.7), Cm(21)
    section.top_margin = section.bottom_margin = Cm(1.5)
    section.left_margin = section.right_margin = Cm(1.5)
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(9)

    document.add_heading(f"Результаты рассылки «{title}»", level=1)
    answered = sum(bool(result["responded"]) for result in results)
    document.add_paragraph(f"Ответили: {answered} из {len(results)}")
    table = document.add_table(rows=1, cols=len(HEADERS))
    table.style = "Table Grid"
    table.autofit = False
    widths = tuple(map(Cm, (2.8, 2.2, 4.2, 3.5, 5, 3.2, 5.8)))
    for cell, value in zip(table.rows[0].cells, HEADERS, strict=True):
        cell.text = value
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for cell, width in zip(table.rows[0].cells, widths, strict=True):
        cell.width = width
    for result in results:
        cells = table.add_row().cells
        for cell, value in zip(cells, _row(result), strict=True):
            cell.text = str(value)
        for cell, width in zip(cells, widths, strict=True):
            cell.width = width

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _row(result: Mapping[str, object]) -> tuple[object, ...]:
    responded_at = result.get("responded_at")
    attachments = result.get("attachments")
    return (
        result["study_group_name"],
        result["vk_user_id"],
        f"{result['last_name']} {result['first_name']}",
        _status(result),
        result.get("text") or "",
        _format_datetime(responded_at),
        "\n".join(_photo_urls(attachments)),
    )


def _status(result: Mapping[str, object]) -> str:
    if not result["responded"]:
        return "Не ответил"
    return "Ответил после дедлайна" if result.get("is_late") else "Ответил"


def _format_datetime(value: object) -> str:
    if not isinstance(value, datetime):
        return ""
    return value.astimezone(MOSCOW).strftime("%d.%m.%Y %H:%M")


def _photo_urls(value: object) -> list[str]:
    if not isinstance(value, list):
        return []
    urls: list[str] = []
    for attachment in value:
        if not isinstance(attachment, dict) or attachment.get("type") != "photo":
            continue
        photo = attachment.get("photo")
        sizes = photo.get("sizes") if isinstance(photo, dict) else None
        if isinstance(sizes, list) and sizes and isinstance(sizes[-1], dict):
            url = sizes[-1].get("url")
            if isinstance(url, str):
                urls.append(url)
    return urls
